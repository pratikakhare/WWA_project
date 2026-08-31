from django.shortcuts import render
from django.http import HttpResponse
from django.views.decorators.http import require_http_methods
from .rfq_cleaner import process_rfq_file
from django.views.decorators.cache import never_cache
from accounts.decorators import role_required

@never_cache
@role_required(["ADMIN", "RFQ"])
def rfq_dashboard(request):
    return render(request, "rfq/rfq_dashboard.html")

@role_required(["ADMIN", "RFQ"])
@require_http_methods(["GET", "POST"])
def rfq_cleaner_view(request):
    context = {}

    if request.method == "POST":
        uploaded_file = request.FILES.get("file")

        if not uploaded_file:
            context["error"] = "Please upload an Excel file."
            return render(request, "rfq/rfq_cleaner.html", context)

        try:
            file_content = process_rfq_file(uploaded_file)

            response = HttpResponse(
                file_content,
                content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            )

            response["Content-Disposition"] = (
                'attachment; filename="cleaned_rfq.xlsx"'
            )

            return response

        except Exception as e:
            context["error"] = str(e)

    return render(request, "rfq/rfq_cleaner.html", context)
# ==============================================================================================


# Excel to Word functionality

import os
from io import BytesIO
from docx.shared import RGBColor
import pandas as pd
from django.conf import settings
from django.shortcuts import render
from django.http import HttpResponse
from docx import Document
from datetime import datetime, timedelta, date
from calendar import monthrange


def remove_none_rows(table):
    """
    Remove entire table rows if any cell contains 'None'.
    """
    rows_to_remove = []
    for row in list(table.rows):
        row_text = " ".join(cell.text.strip() for cell in row.cells)
        if "None" in row_text:
            rows_to_remove.append(row)

    for row in rows_to_remove:
        table._tbl.remove(row._tr)


# --------------------------------------------------
# Load Excel
# --------------------------------------------------
def load_excel():

    excel_path = os.path.join(
        settings.BASE_DIR,
        "complete",
        "data",
        "Working file.xlsx"
    )

    df = pd.read_excel(
        excel_path,
        sheet_name="FAK_Mastersheet",
        header=1
    )

    df.columns = (
        df.columns
        .str.strip()
        .str.upper()
        # .str.replace(r"\s+", " ", regex=True)
        # .str.replace("\n", " ", regex=False)
        # .str.replace(r"[^A-Z0-9 ]", "", regex=True)
    )
    return df
   
    

# --------------------------------------------------
# Home Page
# --------------------------------------------------
@role_required(["ADMIN", "RFQ"])
def index(request):

    df = load_excel()

    customers = (
        df["WWA_CUSTOMER"]
        .dropna()
        .astype(str)
        .sort_values()
        .unique()
        .tolist()
    )

    return render(request, "rfq/excel_word.html", {
    "customers": customers,
    "today": date.today().strftime("%Y-%m-%d"),
    })


# --------------------------------------------------
# Generate Word
# --------------------------------------------------
@role_required(["ADMIN", "RFQ"])
def generate_word(request):

    if request.method != "POST":
        return render(request, "rfq/excel_word.html", {
            "customers": (
                load_excel()["WWA_CUSTOMER"]
                .dropna()
                .astype(str)
                .sort_values()
                .unique()
                .tolist()
            ),
            "today": date.today().strftime("%Y-%m-%d"),
        })

    wwa_customer = request.POST.get("wwa_customer")
    member_deadline_raw = request.POST.get("members_deadline") or request.POST.get("member_deadline")
    
    if not member_deadline_raw:
        return HttpResponse("Members deadline is required.", status=400)

    try:
        member_deadline = datetime.strptime(member_deadline_raw, "%Y-%m-%d").date()
    except ValueError:
        return HttpResponse("Invalid members deadline format.", status=400)
    
    #Calendar Validation: Ensure the member_deadline is not in the past
    if member_deadline < date.today():
        return HttpResponse(
        "Members deadline cannot be a past date.",
        status=400
    )

    validity_half = request.POST.get("validity_half")
    today = date.today()

     #Date Update
    today = date.today()
    if today.month == 12:
        next_month_first = date(today.year + 1, 1, 1)
    else:
        next_month_first = date(today.year, today.month + 1, 1)        
    # Determine the quarter-end month
    quarter_end_month = ((next_month_first.month - 1) // 3 + 1) * 3

    # Last day of the quarter-end month
    last_day = monthrange(next_month_first.year, quarter_end_month)[1]
    quarter_end_date = date(next_month_first.year, quarter_end_month, last_day)

    #Monthly Frequency Logic
    last_day = monthrange(next_month_first.year, next_month_first.month)[1]
    end_of_month = date(next_month_first.year, next_month_first.month, last_day) 
    
    current_month_last_day = monthrange(today.year, today.month)[1]
    end_of_current_month = date(today.year, today.month, current_month_last_day)
    
    # Dropdown For BI-WEEKLY
    if validity_half == "First Half":

        biweekly_effective_date = next_month_first
        biweekly_expiry_date = next_month_first + timedelta(days=14)

    elif validity_half == "Second Half":

        biweekly_effective_date = date(today.year, today.month, 16)
        biweekly_expiry_date = end_of_current_month

    else:

        # Default values
        biweekly_effective_date = next_month_first
        biweekly_expiry_date = next_month_first + timedelta(days=14)
    
    
    
    df = load_excel()
    row = df[df["WWA_CUSTOMER"] == wwa_customer]
    if row.empty:
        return HttpResponse("Customer not found.")
    row = row.iloc[0]
    
    
    # ---------------------------------
    # Year Logic
    # ---------------------------------

    if row["UPDATE_TYPE"] == "OFR & Service Update Request":
        year = row["OFR_VALIDITY_EFFECTIVE_DATE"].year
    else:
        year = row["FOB_PLC_VALIDITY_EFFECTIVE_DATE"].year
     
       
    
    #List Of All Surcharges
    surchanges = ['SOLAS', 'FILING', 'MANDATORY_ICS2_ADMIN_FEE',
       'CONDITIONAL_ICS2_FILING_FEE', 'ORIGIN_E_INF', 'DESTINATION_E_INF',
       'MPCI_FILING_FEE', 'IMO2020', 'GRI', 'ETSS_FEMS', 'ESF', 'E_PSS',
       'LWS_PCS', 'E_BAF', 'WRS']

    
    quarterly_surcharges = []
    monthly_surcharges = []
    bi_weekly_surcharges = []
    
    locked_surcharges = []
    unlocked_surcharges = []

    for surcharge in surchanges:
        value = str(row.get(surcharge, "")).strip().lower()
        if value == "quarterly":
            if str(row.get("FREQUENCY_OFR", "")).strip().lower() == "quarterly":
                locked_surcharges.append(surcharge)
            else:
                unlocked_surcharges.append(surcharge)  
    
    if locked_surcharges:
        quarterly_display = ", ".join(locked_surcharges) + " (Locked)"
    elif unlocked_surcharges:
        quarterly_display = ", ".join(unlocked_surcharges)
    else:
        quarterly_display = "None"            
                  
    # Quarterly date logic
    if locked_surcharges:

        quarterly_effective_date = row["OFR_VALIDITY_EFFECTIVE_DATE"]
        quarterly_expiry_date = row["OFR_VALIDITY_EXPIRY_DATE"]

    elif unlocked_surcharges:

        quarterly_effective_date = next_month_first
        quarterly_expiry_date = quarter_end_date

    else:

        quarterly_effective_date = next_month_first
        quarterly_expiry_date = quarter_end_date              
                  
    for surcharge in surchanges:
        if "Quarterly" in str(row[surcharge]):
            quarterly_surcharges.append(surcharge)
        if "Month" in str(row[surcharge]):
            monthly_surcharges.append(surcharge)
        if  "Weekly" in str(row[surcharge]):
            bi_weekly_surcharges.append(surcharge)
    
    # Monthly display logic with LOCKED word
    if validity_half == "Second Half" and monthly_surcharges:
        monthly_display = ", ".join(monthly_surcharges) + " (LOCKED)"
    elif monthly_surcharges:
        monthly_display = ", ".join(monthly_surcharges)
    else:
        monthly_display = "None"
    
    
    # for surcharge in surchanges:
    #     value = str(row.get(surcharge, "")).strip().lower()
    #     if value.startswith("quarterly"):
    #         quarterly_surcharges.append(surcharge)
    #     if "month" in value:
    #         monthly_surcharges.append(surcharge)
    #     if "weekly" in value:
    #         bi_weekly_surcharges.append(surcharge)
            

    
    # Excluded Surcharges
    excluded = ['SOLAS', 'FILING', 'MANDATORY_ICS2_ADMIN_FEE',
        'CONDITIONAL_ICS2_FILING_FEE', 'ORIGIN_E_INF', 'DESTINATION_E_INF',
        'MPCI_FILING_FEE', 'IMO2020', 'GRI', 'ETSS_FEMS', 'ESF', 'E_PSS',
        'LWS_PCS', 'E_BAF', 'WRS', 'HAZ', '3RD_PARTY_CFS', 'NON_STACKABLE',
        'STORAGE_FREE_DAYS', 'STORAGE_CHARGES', 'BMSB']

    Excluded_surcharges = []
    Excl_as_per_advisory = []
    for exc_surcharge in excluded:
        if "Excluded" in str(row[exc_surcharge]):
            Excluded_surcharges.append(exc_surcharge)
        if "Excl. As Per Advisory" in str(row[exc_surcharge]):
            Excl_as_per_advisory.append(exc_surcharge)        
                
    # ---------------------------------
    # Load Word Template
    # ---------------------------------

    template_path = os.path.join(
        settings.BASE_DIR,
        "complete",
        "data",
        "template.docx"
    )

    doc = Document(template_path)

    # -----------------------------
    # Placeholder Dictionary
    # -----------------------------
    replace_dict = {
        "<<UPDATE TYPE>>": str(row["UPDATE_TYPE"].upper()),
        "<<WWA CUSTOMER>>": str(row["WWA_CUSTOMER"]),
        "<<YEAR>>": str(year),
        "<<HANDLE VIA>>": str(row["HANDLE_VIA"]),
        "<<Commodity>>": str(row["COMMODITY"]),
        "<<Cargo Ratio>>": str(row["CARGO_RATIO"]),
        "<<GLOBAL/REGIONAL ACCOUNT OWNER>>": str(row["GLOBAL_REGIONAL_ACCOUNT_OWNER"]),
        "<<RFQ OWNER>>": str(row["RFQ_OWNER"]),

    #Table Update#

        "<<FOBPLC EFFECTIVE DATE>>": str(row["FOB_PLC_VALIDITY_EFFECTIVE_DATE"].strftime("%d %b %Y")),
        "<<FOBPLC_VALIDITY_EXPIRY_DATE>>": str(row["FOB_PLC_VALIDITY_EXPIRY_DATE"].strftime("%d %b %Y")),
        "<<OFR_VALIDITY_EFFECTIVE_DATE>>": str(row["OFR_VALIDITY_EFFECTIVE_DATE"].strftime("%d %b %Y")),
        "<<OFR_VALIDITY_EXPIRY_DATE>>": str(row["OFR_VALIDITY_EXPIRY_DATE"].strftime("%d %b %Y")),
        "<<OFR_VALIDITY_EFFECTIVE_DATE_QUARTERLY>>": next_month_first.strftime("%d %b %Y"),
        "<<OFR_VALIDITY_EFFECTIVE_DATE_WEEKLY>>": next_month_first.strftime("%d %b %Y"),
        "<<OFFER_VALIDITY_EFFECTIVE_DATE_MONTHLY>>": next_month_first.strftime("%d %b %Y"),
        
        "<<OFR_VALIDITY_EFFECTIVE_DATE_BI_WEEKLY>>": biweekly_effective_date.strftime("%d %b %Y"),
        "<<OFR_VALIDITY_EFFECTIVE_DATE_QUARTERLY>>":quarterly_effective_date.strftime("%d %b %Y"),
        
        
        "<<MONTHLY>>": monthly_display,
        "<<BI-WEEKLY>>": ", ".join(bi_weekly_surcharges) if bi_weekly_surcharges else "None",
        # "<<WEEKLY>>" : ", ".join(weekly_surcharges) if weekly_surcharges else "None",
        "<<OFR_VALIDITY_EFFECTIVE_DATE_WEEKLY>>": str(row["OFR_VALIDITY_EFFECTIVE_DATE"].strftime("%d %b %Y")),
        
        
        "<<FREQUENCY_BIWEEKLY>>": biweekly_expiry_date.strftime("%d %b %Y"),       
        "<<FREQUENCY_WEEKLY>>":(next_month_first + timedelta(days=6)).strftime("%d %b %Y"),
        "<<FREQUENCY_MONTHLY>>": end_of_month.strftime("%d %b %Y"),
        "<<FREQUENCY_QUARTERLY>>":quarterly_expiry_date.strftime("%d %b %Y"),     
        
        #Excluded Surcharges
        "<<excluded>>": "/ ".join(Excluded_surcharges) if Excluded_surcharges else "None",
        "<<EXCL_AS_PER_ADVISORY>>": "/ ".join(Excl_as_per_advisory) if Excl_as_per_advisory else "None",
        
        
        "<<ALTERNATIVE_CFS>>": str(row["ALTERNATIVE_CFS"]),

        "<<MEMBERS_DEADLINE>>": member_deadline.strftime("%d %b %Y")
    }

    # Logic For the (LOCKED) word
    def replace_quarterly(paragraph):

        for run in paragraph.runs:

            if "<<QUARTERLY>>" in run.text:

                # Remove the placeholder from this run
                run.text = run.text.replace("<<QUARTERLY>>", "")

                if locked_surcharges:

                    # Same font/style as template
                    new_run = paragraph.add_run(", ".join(locked_surcharges) + " ")
                    new_run.font.name = run.font.name
                    new_run.font.size = run.font.size
                    new_run.bold = run.bold
                    new_run.italic = run.italic
                    new_run.underline = run.underline

                    # Only "(Locked)" in red
                    lock_run = paragraph.add_run("(LOCKED)")
                    lock_run.font.name = run.font.name
                    lock_run.font.size = run.font.size
                    lock_run.bold = run.bold
                    lock_run.italic = run.italic
                    lock_run.underline = run.underline
                    lock_run.font.color.rgb = RGBColor(255, 0, 0)

                elif unlocked_surcharges:

                    new_run = paragraph.add_run(", ".join(unlocked_surcharges))
                    new_run.font.name = run.font.name
                    new_run.font.size = run.font.size
                    new_run.bold = run.bold
                    new_run.italic = run.italic
                    new_run.underline = run.underline

                else:

                    new_run = paragraph.add_run("None")
                    new_run.font.name = run.font.name
                    new_run.font.size = run.font.size
                    new_run.bold = run.bold
                    new_run.italic = run.italic
                    new_run.underline = run.underline

                break

    
    # Monthly Logic with LOCKED word RED color 

    def replace_monthly(paragraph):

        if "<<MONTHLY>>" not in paragraph.text:
            return

        for run in paragraph.runs:

            if "<<MONTHLY>>" in run.text:

                run.text = run.text.replace("<<MONTHLY>>", "")

                if validity_half == "Second Half" and monthly_surcharges:

                    normal_run = paragraph.add_run(", ".join(monthly_surcharges) + " ")

                    # Copy formatting
                    normal_run.font.name = run.font.name
                    normal_run.font.size = run.font.size
                    normal_run.bold = run.bold
                    normal_run.italic = run.italic
                    normal_run.underline = run.underline

                    locked_run = paragraph.add_run("(LOCKED)")
                    locked_run.font.name = run.font.name
                    locked_run.font.size = run.font.size
                    locked_run.bold = run.bold
                    locked_run.italic = run.italic
                    locked_run.underline = run.underline
                    locked_run.font.color.rgb = RGBColor(255, 0, 0)

                elif monthly_surcharges:

                    normal_run = paragraph.add_run(", ".join(monthly_surcharges))

                    normal_run.font.name = run.font.name
                    normal_run.font.size = run.font.size
                    normal_run.bold = run.bold
                    normal_run.italic = run.italic
                    normal_run.underline = run.underline

                else:

                    paragraph.add_run("None")

                break
    #--------------------------------------------------------
    # Done with Data Filling
    #--------------------------------------------------------
    
    pi_df = pd.read_excel(
        r"C:\Users\pakhare.SHIPCO\complete\complete\complete\data\Customer_Exception.xlsx"
    )
    
    pi_df.columns = (
        pi_df.columns
            .str.strip()
            .str.upper()
            .str.replace(" ", "_")
    )

    pi_mapping = dict(
        zip(
            pi_df["WWA_CUSTOMER"].str.strip(),
            pi_df["PI_LINK"].str.strip()
        )
    )
    
    customer = row["WWA_CUSTOMER"].strip()

    pi_link = pi_mapping.get(customer, "")
    
    
    
    
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.opc.constants import RELATIONSHIP_TYPE


    def add_hyperlink(paragraph, url, text):
        """
        Add a clickable hyperlink to a Word paragraph.
        """

        # Create relationship
        part = paragraph.part
        r_id = part.relate_to(
            url,
            RELATIONSHIP_TYPE.HYPERLINK,
            is_external=True
        )

        # Create hyperlink element
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), r_id)

        # Create run
        new_run = OxmlElement("w:r")

        # Style (blue + underline)
        rPr = OxmlElement("w:rPr")

        color = OxmlElement("w:color")
        color.set(qn("w:val"), "0000FF")
        rPr.append(color)

        underline = OxmlElement("w:u")
        underline.set(qn("w:val"), "single")
        rPr.append(underline)
        
        # 9 pt = 18 half-points
        size = OxmlElement("w:sz")
        size.set(qn("w:val"), "18")
        rPr.append(size)

        size_cs = OxmlElement("w:szCs")
        size_cs.set(qn("w:val"), "18")
        rPr.append(size_cs)

        new_run.append(rPr)

        # Text
        text_elem = OxmlElement("w:t")
        text_elem.text = text
        new_run.append(text_elem)

        hyperlink.append(new_run)

        paragraph._p.append(hyperlink)
        
    # Replace the LINK Word    
    for paragraph in doc.paragraphs:

        if "<<LINK>>" in paragraph.text:

            paragraph.text = ""

            if pi_link:
                add_hyperlink(
                    paragraph,
                    pi_link,
                    "Click here to view PI"
                )    
            
        
        
        
        
        
        
    # ---------------------------------
    # Replace Function
    # ---------------------------------

    def replace_in_paragraph(paragraph):
        for key, value in replace_dict.items():
            if key in paragraph.text:
                for run in paragraph.runs:
                    run.text = run.text.replace(key, str(value))

    # ---------------------------------
    # Replace in Normal Paragraphs
    # ---------------------------------

    for paragraph in doc.paragraphs:
        replace_quarterly(paragraph)
        replace_monthly(paragraph)
        replace_in_paragraph(paragraph)

    # ---------------------------------
    # Replace in Tables
    # ---------------------------------

    for table in doc.tables:
        for table_row in table.rows:
            for cell in table_row.cells:
                for paragraph in cell.paragraphs:
                    replace_quarterly(paragraph)
                    replace_monthly(paragraph)
                    replace_in_paragraph(paragraph)
                    

        remove_none_rows(table)

    # ---------------------------------
    # Gray Out Alternative Ports Section
    # ---------------------------------

    alternative_cfs = str(row["ALTERNATIVE_CFS"]).strip().lower()
    if alternative_cfs == "not allowed":
        gray_section = False
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            # Start graying from this heading
            if text.startswith("Where alternative ports are quoted"):
                gray_section = True

            # Stop BEFORE this heading
            if text.startswith("Service Information Requirements"):
                break

            if gray_section:
                for run in paragraph.runs:
                    run.font.color.rgb = RGBColor(128, 128, 128)

    # ---------------------------------
    # Download Word
    # ---------------------------------

    output = BytesIO()

    doc.save(output)

    output.seek(0)

    response = HttpResponse(
        output.getvalue(),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

    response["Content-Disposition"] = (
        f'attachment; filename="{wwa_customer}.docx"'
    )
    return response







from django.http import JsonResponse

@role_required(["ADMIN", "RFQ"])
def customer_details(request):

    customer = request.GET.get("customer")

    df = load_excel()

    row = df[df["WWA_CUSTOMER"] == customer]

    if row.empty:
        return JsonResponse({"success": False})

    row = row.iloc[0]

    surchanges = [
        'SOLAS', 'FILING', 'MANDATORY_ICS2_ADMIN_FEE',
        'CONDITIONAL_ICS2_FILING_FEE', 'ORIGIN_E_INF',
        'DESTINATION_E_INF', 'MPCI_FILING_FEE',
        'IMO2020', 'GRI', 'ETSS_FEMS', 'ESF',
        'E_PSS', 'LWS_PCS', 'E_BAF', 'WRS',
        'HAZ', '3RD_PARTY_CFS', 'NON_STACKABLE',
        'STORAGE_FREE_DAYS', 'STORAGE_CHARGES',
        'BMSB'
    ]

    quarterly_surcharges = []
    monthly_surcharges = []
    bi_weekly_surcharges = []

    for surcharge in surchanges:

        value = str(row.get(surcharge, "")).strip()

        if "Quarterly" in value:
            quarterly_surcharges.append(surcharge)

        if "Month" in value:
            monthly_surcharges.append(surcharge)

        if "Bi-Weekly" in value:
            bi_weekly_surcharges.append(surcharge)

    return JsonResponse({

        "success": True,

        "Quarterly": ", ".join(quarterly_surcharges) if quarterly_surcharges else "No Surcharges Available",

        "Monthly": ", ".join(monthly_surcharges) if monthly_surcharges else "No Surcharges Available",

        "BI_WEEKLY": ", ".join(bi_weekly_surcharges) if bi_weekly_surcharges else "No Surcharges Available",

        "update_type": row["UPDATE_TYPE"],
        "handle_via": row["HANDLE_VIA"],
        "commodity": row["COMMODITY"],
        "cargo_ratio": row["CARGO_RATIO"],
        "frequency_ofr": row["FREQUENCY_OFR"],
        "currency_ofr": row["CURRENCY_OFR"],
        "rfq_owner": row["RFQ_OWNER"],
    })